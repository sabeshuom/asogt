#!/usr/bin/env python
# -*- coding: utf-8 -*-
from __future__ import unicode_literals

import os
import sys
import requests
import json
import xlsxwriter
import re
import xlrd
import string
import numpy as np
from docx import Document
from docx.shared import Inches
import time
import pandas as pd

sys.path.append("../../")
from asogt.settings import MEDIA_ROOT, BASE_DIR
from core.data_utils import init_sess,\
    process_results_for_seating_number,\
    get_results,\
    get_ref_data_from_excel,\
    DIVISION_ORDER,\
    sort_std_keys_for_division

from core.unicode_to_bamini import unicode2bamini

from core.write_timetable import add_time_table

BOOK_GRADES = ["First Prize",
               "Second Prize",
               "Third Prize",
               "Grade A",
               "Grade B",
               "Grade C",
               "Participated"]

STATE_TAMIL = {"QLD": "Fapd;];yhe;J",
               "NSW":	"epA+ rTj; Nty;];",
               "VIC": "tpf;Nuhupah",
               "SA": "njw;F M];jpNuypah",
               "ACT": "fd;nguh",
               "WA": "Nkw;F M];jpNuypah",
               "NZW": "ntypq;ld;> epA+rpyhe;J",
               "NZH": "`hkpy;ld;> epA+rpyhe;J"}


GRADE_INFO = {"First Prize": (" ", "Kjw; ghpR"),
              "Second Prize": (" ", ",uz;lhk; ghpR"),
              "Third Prize":  (" ", "%d;whk; ghpR"),
              "Grade A": ("A ", "epiy"),
              "Grade B": ("B ", "epiy"),
              "Grade C": ("C ", "epiy"),
              "Participated": (" ", "gq;Fgw;wpdhu;"),
              }


def get_updated_student_data_map(student_data_map):
    
    df = pd.ExcelFile("map.xlsx")
    new_map = df.parse(
        "Student List").fillna('').set_index("Student Number").transpose()
    
    for std_no in student_data_map:
        student_data_map[std_no].seat_pos = new_map[std_no]["Seating Number"]

    return student_data_map



def export_to_excel(xls_wb, state,  year, exam_category, username, password):

    sess = init_sess(username, password)
    results = get_results(sess, state=state, year=year,
                          competition="All", exam_category=exam_category)
    ordered_results, division_comp_map, student_data_map = process_results_for_seating_number(
        results)
    student_data_map =  get_updated_student_data_map(student_data_map)

    wb = xlsxwriter.Workbook(xls_wb)
    row_height = 25
    row_title_height = 35

    div_header_format = wb.add_format({
        'font_name': "Bamini",
        'align': 'left',
        'bg_color': '#ebf0df',
        'bold': 1,
        'border': 2,
        'font_size': 14,
        'valign': 'vcenter'})

    bamini_cell_format = wb.add_format({
        'font_name': "Bamini",
        'align': 'left',
        'font_size': 14,
        'border': 1,
        'valign': 'vcenter'})

    uc_cell_format = wb.add_format({
        'font_name': "Calibri (Body)",
        'align': 'left',
        'font_size': 14,
        'border': 1,
        'valign': 'vcenter'})

    for division, division_prefix in DIVISION_ORDER:
        ws = wb.add_worksheet(division)
        ws.set_default_row(row_height)
        comps = [comp for comp in sorted(
            division_comp_map[division], key=lambda x: division_comp_map[division][x], reverse=True)]
        comps_bamini = [unicode2bamini(comp) for comp in comps]
        div_header = ["Seat POS", "khztu; ,y.",
                      "KOg; ngau;", "KOg; ngau;"] + comps_bamini
        ws.write_row(0, 0, div_header, div_header_format)
        ws.set_row(0, row_title_height)

        division_data = ordered_results[division]
        header_r = 1
        if division_prefix == "G":
            sorted_stds = sort_std_keys_for_division(division_data)
        else:
            sorted_stds = sorted(division_data, key=lambda x: int(
                student_data_map[x].seat_pos[-3:]))

        for r, std_no in enumerate(sorted_stds):
            std_data = student_data_map[std_no]
            name_uc = std_data.name_t
            seat_pos = std_data.seat_pos
            std_comps = division_data[std_no]
            name_bamini = std_data.name_bamini
            c = 0
            ws.write_string(header_r + r, c, std_no, uc_cell_format)
            c += 1
            ws.write_string(header_r + r, c, seat_pos, uc_cell_format)
            c += 1
            ws.write_string(header_r + r, c, name_uc, uc_cell_format)
            c += 1
            ws.write_string(header_r + r, c, name_bamini, bamini_cell_format)

            for comp in comps:
                c += 1
                ws.write_rich_string(header_r + r, c, " ", uc_cell_format)
                if comp in std_comps:
                    grade = std_comps[comp]
                    grade_e, grade_bamini = GRADE_INFO[grade]
                    if grade_e != "":
                        ws.write_rich_string(header_r + r, c,
                                             uc_cell_format, grade_e, bamini_cell_format, grade_bamini, uc_cell_format)
                    else:
                        ws.write_rich_string(
                            header_r + r, c, bamini_cell_format, grade_bamini, uc_cell_format)

        ws.set_column('A:A', 15)
        ws.set_column('B:B', 45)
        ws.set_column('C:C', 35)
        ws.set_column('D:G', 45)

    wb.close()


# class Logger(object)

#     def __init__(self, log_str=""):
#         self.log_str = log_str

#     def add(log_str):
#         self.


def export_to_docx(word_doc, state,  year,  exam_category, username, password):
    sess = init_sess(username, password)
    results = get_results(sess, state=state, year=year,
                          competition="All", exam_category=exam_category)

    ordered_results, division_comp_map, student_data_map = process_results_for_seating_number(
        results)
    student_data_map =  get_updated_student_data_map(student_data_map)
    template = os.path.join(MEDIA_ROOT, "book_template.docx")
    document = Document(template)

    for division, division_prefix in DIVISION_ORDER:
        comps = [comp for comp in sorted(
            division_comp_map[division], key=lambda x: division_comp_map[division][x], reverse=True)]
        comps_bamini = [unicode2bamini(comp) for comp in comps]
        # convert division unicode to bamini
        division_bamini = unicode2bamini(division)
        # division heading with state info
        division_heading = "{:s} - {:s} -  ghpRngw;Nwhh; gl;bay;".format(
            division_bamini, STATE_TAMIL[state])
        # Title
        document.add_paragraph(division_heading, style='Section Title Tamil')

        # add table for each division
        table = document.add_table(
            rows=1, cols=len(comps) + 2, style='Table Grid')
        hdr_cells = table.rows[0].cells

        # write headers
        # First column Aasana ilakkam
        hdr_cells[0].text = "Mrd ,y."
        hdr_cells[0].paragraphs[0].style = document.styles["Table Header Tamil"]

        # first column = fullname in bamini
        hdr_cells[1].text = "KOg;ngah;"
        hdr_cells[1].paragraphs[0].style = document.styles["Table Header Tamil"]

        # next columns with competitions details in bamini
        for c, comp_bamini in enumerate(comps_bamini):
            cell = hdr_cells[c+2]
            cell.text = comp_bamini
            cell.paragraphs[0].style = document.styles["Table Header Tamil"]

        division_data = ordered_results[division]

        if division_prefix == "G":
            sorted_stds = sort_std_keys_for_division(division_data)
        else:
            sorted_stds = sorted(division_data, key=lambda x: int(
                student_data_map[x].seat_pos[-3:]))

        for std_no in sorted_stds:
            std_data = student_data_map[std_no]
            name_uc = std_data.name_t
            seat_pos = std_data.seat_pos
            name_bamini = std_data.name_bamini
            std_comps = division_data[std_no]

            row_cells = table.add_row().cells
            row_cells[0].text = seat_pos
            row_cells[0].paragraphs[0].style = document.styles["Table Cell Eng"]
            row_cells[1].text = name_bamini
            row_cells[1].paragraphs[0].style = document.styles["Table Cell Tamil Left"]

            # loop through sored order on comps
            for c, comp in enumerate(comps):
                cell = row_cells[c+2]
                cell.text = ""
                if comp in std_comps:
                    grade = std_comps[comp]
                    grade_e, grade_bamini = GRADE_INFO[grade]
                    if grade_e != " ":
                        eng = cell.paragraphs[0].add_run(grade_e + " ")
                        eng.font.name = "Calibri"
                    tamil = cell.paragraphs[0].add_run(grade_bamini)
                    tamil.font.name = "Bamini"

    # document.add_page_break()
    document.save(word_doc)


if __name__ == "__main__":
    username = "sabesan"
    password = "Sabesan4NSW"
    state = "NSW"
    year = "2018"
    exam_category = ["State", "Final"]
    xls_wb = "book.xlsx"
    word_doc = "book.docx"
    export_to_excel(xls_wb, state,  year, exam_category, username, password)
    export_to_docx(word_doc, state,  year, exam_category,  username, password)
